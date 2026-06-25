"""Ingest + index a principal's context documents, incrementally.

Design goals (per product spec):
  * Drop files in ``data/principal_docs/<slug>/`` and index them.
  * INCREMENTAL: a file is only re-read and re-sent to the LLM when it is new or
    its content hash changed. Unchanged files are skipped — we never re-process
    the whole corpus.
  * The LLM (Claude via the insight provider) distills each doc into a summary,
    verbatim proof points, and retrieval themes. These ground relevance research
    and outreach so we can relate the principal's experience to a prospect's.
"""
from __future__ import annotations

import hashlib
import logging
import re
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional

from sqlalchemy import select
from sqlalchemy.orm import Session

from app.models.principal import Principal
from app.models.principal_document import PrincipalDocument
from app.services.insights import get_insight_provider

logger = logging.getLogger(__name__)

# backend/data/principal_docs
DEFAULT_DOCS_ROOT = Path(__file__).resolve().parents[3] / "data" / "principal_docs"

SUPPORTED_SUFFIXES = {".pdf", ".docx", ".txt", ".md"}

# Relevance thresholds (0–100) for document status + outreach inclusion.
RELEVANCE_CORE = 55.0       # status: indexed — used fully in dossier/outreach
RELEVANCE_PERIPHERAL = 35.0  # status: peripheral — partial use, flagged in UI
# Below RELEVANCE_PERIPHERAL → status: irrelevant — stored but excluded from outreach

_PROCESSED_STATUSES = frozenset({"indexed", "peripheral", "irrelevant"})


def _slugify(name: str) -> str:
    return re.sub(r"[^a-z0-9]+", "-", (name or "").lower()).strip("-") or "principal"


def _extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    if suffix == ".docx":
        import docx  # python-docx

        document = docx.Document(str(path))
        return "\n".join(p.text for p in document.paragraphs)
    raise ValueError(f"Unsupported file type: {suffix}")


def _hash_file(path: Path) -> str:
    h = hashlib.sha256()
    h.update(path.read_bytes())
    return h.hexdigest()


def principal_docs_dir(principal_name: str, root: Optional[Path] = None) -> Path:
    """Resolve the drop folder for a principal, tolerant of naming.

    Tries the full-name slug ("dalbir-bains") then the first-name slug ("dalbir")
    so files land correctly whichever folder the user created. Falls back to the
    full-name slug path when neither exists yet.
    """
    base = root or DEFAULT_DOCS_ROOT
    candidates = [_slugify(principal_name)]
    first = (principal_name or "").strip().split()
    if first:
        candidates.append(_slugify(first[0]))
    for slug in candidates:
        path = base / slug
        if path.exists():
            return path
    return base / candidates[0]


def _principal_context(principal: Principal) -> dict:
    return {
        "name": principal.name,
        "headline": principal.headline,
        "document_focus": principal.document_focus,
        "bio": principal.bio,
        "background": principal.background,
        "focus_areas": principal.focus_areas,
        "target_sectors": principal.target_sectors,
        "investment_themes": principal.investment_themes,
        "value_props": principal.value_props,
        "opportunity_types": principal.opportunity_types,
        "geographies": principal.geographies,
    }


def _status_for_score(score: float) -> str:
    if score >= RELEVANCE_CORE:
        return "indexed"
    if score >= RELEVANCE_PERIPHERAL:
        return "peripheral"
    return "irrelevant"


def _process_one_file(
    db: Session,
    *,
    principal_id: int,
    path: Path,
    record: Optional[PrincipalDocument],
    provider,
    p_ctx: dict,
    force: bool,
) -> tuple[Optional[PrincipalDocument], dict]:
    """Index a single file. Returns (record, result dict for API/UI)."""
    file_hash = _hash_file(path)
    unchanged = (
        record is not None
        and record.file_hash == file_hash
        and record.status in _PROCESSED_STATUSES
    )
    if unchanged and not force:
        return record, {"file": path.name, "action": "skipped", "relevance_score": record.relevance_score}

    try:
        text = _extract_text(path)
        index = provider.index_document(text=text, filename=path.name, principal=p_ctx)
    except Exception as exc:  # noqa: BLE001
        logger.warning("Failed to index %s: %s", path.name, exc)
        return record, {"file": path.name, "action": "failed", "error": str(exc)}

    is_new = record is None
    if record is None:
        record = PrincipalDocument(principal_id=principal_id, filename=path.name)
        db.add(record)
    record.file_hash = file_hash
    record.raw_text = text
    record.char_count = len(text)
    record.summary = index.summary
    record.key_facts = index.key_facts if index.relevance_score >= RELEVANCE_PERIPHERAL else []
    record.themes = index.themes if index.relevance_score >= RELEVANCE_PERIPHERAL else []
    record.doc_type = index.doc_type
    record.relevance_score = index.relevance_score
    record.relevance_note = index.relevance_note
    record.status = _status_for_score(index.relevance_score)
    record.indexed_by = index.generated_by
    record.indexed_at = datetime.utcnow()
    db.flush()

    action = record.status if record.status != "indexed" else "indexed"
    result = {
        "file": path.name,
        "action": action,
        "relevance_score": index.relevance_score,
        "status": record.status,
        "proof_points": len(record.key_facts or []),
    }
    if is_new:
        result["is_new"] = True
    else:
        result["is_new"] = False
    return record, result


def ingest_single_document(
    db: Session,
    principal_id: int,
    principal_name: str,
    filename: str,
    *,
    root: Optional[Path] = None,
    force: bool = False,
) -> Dict[str, object]:
    """Index one document by filename. Used for per-file progress in the UI."""
    folder = principal_docs_dir(principal_name, root)
    path = folder / filename
    if not path.is_file():
        return {"file": filename, "action": "failed", "error": "File not found on disk"}

    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        reason = (
            "Legacy .doc format is not supported — open in Word and Save As .docx"
            if suffix == ".doc"
            else f"Unsupported type: {suffix}"
        )
        return {"file": filename, "action": "rejected", "error": reason}

    record = db.execute(
        select(PrincipalDocument).where(
            PrincipalDocument.principal_id == principal_id,
            PrincipalDocument.filename == filename,
        )
    ).scalar_one_or_none()

    provider = get_insight_provider()
    principal_row = db.get(Principal, principal_id)
    p_ctx = _principal_context(principal_row) if principal_row else {"name": principal_name}

    _, result = _process_one_file(
        db,
        principal_id=principal_id,
        path=path,
        record=record,
        provider=provider,
        p_ctx=p_ctx,
        force=force,
    )
    db.commit()
    return result


def ingest_principal_docs(
    db: Session,
    principal_id: int,
    principal_name: str,
    *,
    root: Optional[Path] = None,
    force: bool = False,
) -> Dict[str, object]:
    """Index new/changed docs for a principal. Returns a per-run summary.

    ``force=True`` re-indexes every file even if unchanged (e.g. after changing
    the extraction prompt). Otherwise only new or modified files are processed.
    """
    folder = principal_docs_dir(principal_name, root)
    summary: Dict[str, object] = {
        "folder": str(folder),
        "indexed": 0,
        "updated": 0,
        "skipped": 0,
        "failed": 0,
        "files": [],
    }
    if not folder.exists():
        summary["error"] = f"Folder not found: {folder}"
        return summary

    existing = {
        d.filename: d
        for d in db.execute(
            select(PrincipalDocument).where(PrincipalDocument.principal_id == principal_id)
        ).scalars().all()
    }
    provider = get_insight_provider()
    principal_row = db.get(Principal, principal_id)
    p_ctx = _principal_context(principal_row) if principal_row else {"name": principal_name}

    for path in sorted(folder.iterdir()):
        if not path.is_file() or path.suffix.lower() not in SUPPORTED_SUFFIXES:
            continue
        if path.name.lower() == "readme.md":
            continue

        record = existing.get(path.name)
        _, result = _process_one_file(
            db,
            principal_id=principal_id,
            path=path,
            record=record,
            provider=provider,
            p_ctx=p_ctx,
            force=force,
        )
        action = result.get("action")
        if action == "skipped":
            summary["skipped"] = int(summary["skipped"]) + 1  # type: ignore[index]
        elif action == "failed":
            summary["failed"] = int(summary["failed"]) + 1  # type: ignore[index]
        elif result.get("is_new"):
            summary["indexed"] = int(summary["indexed"]) + 1  # type: ignore[index]
        elif action in ("indexed", "peripheral", "irrelevant"):
            summary["updated"] = int(summary["updated"]) + 1  # type: ignore[index]
        summary["files"].append(result)  # type: ignore[union-attr]

    db.commit()
    return summary


# --- Retrieval / dossier ---------------------------------------------------


def _all_docs(db: Session, principal_id: int, *, for_outreach: bool = True) -> List[PrincipalDocument]:
    """Docs usable for dossier/outreach. Excludes irrelevant uploads."""
    stmt = select(PrincipalDocument).where(PrincipalDocument.principal_id == principal_id)
    if for_outreach:
        stmt = stmt.where(PrincipalDocument.status.in_(("indexed", "peripheral")))
    return db.execute(stmt).scalars().all()


def build_dossier_summary(db: Session, principal_id: int, *, max_facts: int = 40) -> dict:
    """Structured dossier for the Principals UI — shows what indexing extracted."""
    all_docs = db.execute(
        select(PrincipalDocument).where(PrincipalDocument.principal_id == principal_id)
    ).scalars().all()
    usable = [d for d in all_docs if d.status in ("indexed", "peripheral")]

    facts: List[str] = []
    themes: set[str] = set()
    for d in usable:
        cap = 3 if d.status == "peripheral" else 12
        for f in (d.key_facts or [])[:cap]:
            if f and f not in facts:
                facts.append(f)
        for t in d.themes or []:
            themes.add(t)
    facts = facts[:max_facts]

    documents = []
    for d in sorted(all_docs, key=lambda x: (x.relevance_score or 0), reverse=True):
        documents.append({
            "id": d.id,
            "filename": d.filename,
            "doc_type": d.doc_type,
            "status": d.status,
            "relevance_score": d.relevance_score,
            "relevance_note": d.relevance_note,
            "summary": d.summary,
            "key_facts": d.key_facts or [],
            "themes": d.themes or [],
            "indexed_at": d.indexed_at.isoformat() if d.indexed_at else None,
        })

    return {
        "documents_total": len(all_docs),
        "documents_usable": len(usable),
        "proof_points_total": sum(len(d.key_facts or []) for d in usable),
        "proof_points_unique": len(facts),
        "themes": sorted(themes),
        "top_proof_points": facts[:15],
        "documents": documents,
        "used_in": [
            "Discovery — scopes who to find (titles, sectors, geographies from profile)",
            "Research — Claude reads prospect + injects your proof points for fit scoring",
            "Outreach — picks the 1–2 proof points most relevant to each person for the email",
        ],
    }


def build_principal_dossier(db: Session, principal_id: int, *, max_facts: int = 30) -> str:
    """Concatenate indexed proof points + themes into a compact context block."""
    docs = _all_docs(db, principal_id)
    if not docs:
        return ""
    facts: List[str] = []
    themes: set[str] = set()
    for d in docs:
        if d.status == "peripheral":
            # Use peripheral docs but cap contribution so core docs dominate.
            cap = 3
        else:
            cap = 8
        for f in (d.key_facts or [])[:cap]:
            if f and f not in facts:
                facts.append(f)
        for t in d.themes or []:
            themes.add(t)
    facts = facts[:max_facts]
    parts = []
    if facts:
        parts.append("PROOF POINTS (verbatim from the principal's documents):\n- " + "\n- ".join(facts))
    if themes:
        parts.append("EXPERTISE THEMES: " + ", ".join(sorted(themes)))
    return "\n\n".join(parts)


def filter_usable_proof_points(
    facts: List[str], *, max_items: int = 2, max_len: int = 100
) -> List[str]:
    """Drop resume-index noise and over-long lines unsuitable for email copy."""
    junk = re.compile(
        r"target roles identified|president,\s*ceo,\s*coo,\s*cfo|"
        r"private equity operating partner,\s*managing partner",
        re.IGNORECASE,
    )
    out: List[str] = []
    for raw in facts or []:
        fact = (raw or "").strip()
        if not fact or junk.search(fact):
            continue
        if len(fact) > max_len:
            fact = fact[: max_len - 1].rsplit(" ", 1)[0] + "…"
        if fact not in out:
            out.append(fact)
        if len(out) >= max_items:
            break
    return out


def retrieve_relevant_proof_points(
    db: Session, principal_id: int, query: str, *, k: int = 6
) -> List[str]:
    """Lightweight keyword-overlap retrieval of the proof points most relevant
    to a prospect (so outreach can relate shared expertise without re-reading
    the whole corpus)."""
    docs = _all_docs(db, principal_id)
    if not docs:
        return []
    q_tokens = set(re.findall(r"[a-z0-9]+", (query or "").lower()))
    if not q_tokens:
        # No query signal: return a few high-signal facts as a fallback.
        flat = [f for d in docs for f in (d.key_facts or [])]
        return filter_usable_proof_points(flat, max_items=k, max_len=100)

    scored: List[tuple[float, str]] = []
    seen: set[str] = set()
    for d in docs:
        doc_weight = 1.0 if d.status == "indexed" else 0.45
        theme_tokens = set(
            t for theme in (d.themes or []) for t in re.findall(r"[a-z0-9]+", theme.lower())
        )
        for fact in d.key_facts or []:
            if fact in seen:
                continue
            seen.add(fact)
            f_tokens = set(re.findall(r"[a-z0-9]+", fact.lower()))
            overlap = (
                len(q_tokens & f_tokens) + 0.5 * len(q_tokens & theme_tokens)
            ) * doc_weight
            if overlap > 0:
                scored.append((overlap, fact))
    scored.sort(key=lambda t: t[0], reverse=True)
    return filter_usable_proof_points([f for _, f in scored], max_items=k, max_len=100)
